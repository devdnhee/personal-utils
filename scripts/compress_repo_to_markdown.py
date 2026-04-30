#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = [
#   "fire",
#   "rich",
#   "python-docx",
# ]
# ///
"""Compress a git repository to a single markdown (or docx) file, and decompress it back."""

import re
import subprocess
from pathlib import Path

import fire
from rich.console import Console

console = Console()


def _get_git_files(repo_path: Path) -> list[Path]:
    result = subprocess.run(
        ["git", "ls-files"],
        capture_output=True,
        text=True,
        check=True,
        cwd=repo_path,
    )
    return [repo_path / f for f in result.stdout.splitlines() if f]


def _build_tree(files: list[Path], repo_path: Path) -> dict:
    tree: dict = {}
    for f in files:
        rel = f.relative_to(repo_path)
        parts = rel.parts
        node = tree
        for part in parts[:-1]:
            node = node.setdefault(part, {})
        node[parts[-1]] = f
    return tree


def _render_tree(node: "dict | Path", name: str, depth: int, lines: list[str]) -> None:
    heading = "#" * depth
    lines.append(f"{heading} {name}")
    lines.append("")
    if isinstance(node, Path):
        lang = node.suffix.lstrip(".") or ""
        try:
            content = node.read_text(encoding="utf-8", errors="replace")
        except Exception:
            content = ""
        # Use enough backticks to not clash with content
        fence = "```"
        while f"\n{fence}" in f"\n{content}":
            fence += "`"
        lines.append(f"{fence}{lang}")
        lines.append(content)
        lines.append(fence)
        lines.append("")
    else:
        for key in sorted(node.keys()):
            _render_tree(node[key], key, depth + 1, lines)


def encode(repo: str, output: str = "output.md") -> None:
    """Compress a git repository to a markdown or docx file.

    Args:
        repo: Local path to the git repository.
        output: Output file path (default: output.md). Use .docx for Word format.
    """
    repo_path = Path(repo).resolve()
    output_path = Path(output)
    to_docx = output_path.suffix.lower() == ".docx"

    files = _get_git_files(repo_path)
    tree = _build_tree(files, repo_path)

    lines: list[str] = [f"# {repo_path.name}", ""]
    for key in sorted(tree.keys()):
        _render_tree(tree[key], key, 2, lines)

    md_content = "\n".join(lines)

    if to_docx:
        from docx import Document  # type: ignore

        doc = Document()
        heading_re = re.compile(r"^(#{1,6})\s+(.*)")
        for line in lines:
            m = heading_re.match(line)
            if m:
                level = len(m.group(1))
                doc.add_heading(m.group(2), level=min(level, 9))
            else:
                doc.add_paragraph(line)
        doc.save(output_path)
        console.log(f"Encoded {len(files)} files to [bold]{output_path}[/bold]")
    else:
        output_path.with_suffix(".md") if to_docx else output_path
        md_path = output_path
        md_path.write_text(md_content, encoding="utf-8")
        console.log(f"Encoded {len(files)} files to [bold]{md_path}[/bold]")


_MODES = ("overwrite", "append", "ignore", "block")


def _parse_markdown(md_content: str) -> dict[str, str]:
    """Return {relative_path: file_content} from markdown produced by encode."""
    lines = md_content.splitlines()
    heading_re = re.compile(r"^(#{1,6})\s+(.*)")
    fence_re = re.compile(r"^(`{3,})")

    path_stack: list[tuple[int, str]] = []  # (depth, name)
    file_map: dict[str, str] = {}
    i = 0

    while i < len(lines):
        m = heading_re.match(lines[i])
        if not m:
            i += 1
            continue

        depth = len(m.group(1))
        name = m.group(2).strip()

        # Trim stack to parents of current depth
        path_stack = [(d, n) for d, n in path_stack if d < depth]
        path_stack.append((depth, name))

        if depth == 1:
            # Repo root heading — skip
            i += 1
            continue

        # Look ahead past blank lines to see what follows
        j = i + 1
        while j < len(lines) and lines[j].strip() == "":
            j += 1

        if j >= len(lines) or heading_re.match(lines[j]):
            # Directory node — no content
            i += 1
            continue

        fm = fence_re.match(lines[j]) if j < len(lines) else None
        if fm:
            # File node — content is in a fenced block
            opening_fence = fm.group(1)
            content_lines: list[str] = []
            k = j + 1
            while k < len(lines):
                if lines[k].strip() == opening_fence:
                    break
                content_lines.append(lines[k])
                k += 1
            # Strip only leading blank lines (preserve trailing newline from original)
            while content_lines and content_lines[0] == "":
                content_lines.pop(0)

            rel_path = "/".join(n for _, n in path_stack[1:])  # skip repo heading
            file_map[rel_path] = "\n".join(content_lines)
            i = k + 1
        else:
            # Unexpected format — skip
            i += 1

    return file_map


def decode(input: str, output: str = ".", mode: str = "block") -> None:
    """Decompress a markdown or docx file back to a directory structure.

    Args:
        input: Input markdown or docx file.
        output: Output directory (default: current directory).
        mode: Conflict resolution mode — overwrite | append | ignore | block.
    """
    if mode not in _MODES:
        raise ValueError(f"--mode must be one of {_MODES}, got '{mode}'")

    input_path = Path(input)

    if input_path.suffix.lower() == ".docx":
        from docx import Document  # type: ignore

        doc = Document(input_path)
        raw_lines: list[str] = []
        for para in doc.paragraphs:
            style = para.style.name
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                raw_lines.append("#" * level + " " + para.text)
            else:
                raw_lines.append(para.text)
        md_content = "\n".join(raw_lines)
    else:
        md_content = input_path.read_text(encoding="utf-8")

    h1_match = re.search(r"^# (.+)$", md_content, re.MULTILINE)
    if not h1_match:
        raise ValueError("No H1 heading found in input — cannot determine repo name.")
    repo_name = h1_match.group(1).strip()

    file_map = _parse_markdown(md_content)
    output_dir = Path(output).resolve() / repo_name
    output_dir.mkdir(parents=True, exist_ok=True)

    for rel_path, content in file_map.items():
        dest = output_dir / rel_path
        dest.parent.mkdir(parents=True, exist_ok=True)

        if dest.exists():
            if mode == "block":
                raise FileExistsError(f"[block] Conflict: {dest} already exists.")
            elif mode == "ignore":
                console.log(f"[yellow]ignore[/yellow]    {rel_path}")
                continue
            elif mode == "overwrite":
                dest.write_text(content, encoding="utf-8")
                console.log(f"[red]overwrite[/red]  {rel_path}")
            elif mode == "append":
                existing = dest.read_text(encoding="utf-8")
                dest.write_text(existing + "\n----\n" + content, encoding="utf-8")
                console.log(f"[blue]append[/blue]     {rel_path}")
        else:
            dest.write_text(content, encoding="utf-8")
            console.log(f"[green]create[/green]     {rel_path}")

    console.log(f"Decoded {len(file_map)} files to [bold]{output_dir}[/bold]")


if __name__ == "__main__":
    fire.Fire({"encode": encode, "decode": decode})
