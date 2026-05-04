#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = [
#   "rich",
#   "python-docx",
# ]
# ///
"""Compress a git repository to a single markdown (or docx) file, and decompress it back."""

import argparse
import difflib
import re
import subprocess
from pathlib import Path

from rich.console import Console

console = Console()


def _get_git_files(repo_path: Path) -> list[Path]:
    result = subprocess.run(
        ["git", "ls-files"],
        capture_output=True,
        text=True,
        check=True,
        cwd=repo_path,
    )
    return [repo_path / f for f in result.stdout.splitlines() if f]


def _build_tree(files: list[Path], repo_path: Path) -> dict:
    tree: dict = {}
    for f in files:
        rel = f.relative_to(repo_path)
        parts = rel.parts
        node = tree
        for part in parts[:-1]:
            node = node.setdefault(part, {})
        node[parts[-1]] = f
    return tree


def _render_tree(node: "dict | Path", name: str, depth: int, lines: list[str]) -> None:
    heading = "#" * depth
    lines.append(f"{heading} {name}")
    lines.append("")
    if isinstance(node, Path):
        lang = node.suffix.lstrip(".") or ""
        try:
            content = node.read_text(encoding="utf-8", errors="replace")
        except Exception:
            content = ""
        fence = "```"
        while f"\n{fence}" in f"\n{content}":
            fence += "`"
        lines.append(f"{fence}{lang}")
        lines.append(content)
        lines.append(fence)
        lines.append("")
    else:
        for key in sorted(node.keys()):
            _render_tree(node[key], key, depth + 1, lines)


def encode(args: argparse.Namespace) -> None:
    repo_path = Path(args.repo).resolve()
    output_path = Path(args.output)
    to_docx = output_path.suffix.lower() == ".docx"

    files = _get_git_files(repo_path)
    tree = _build_tree(files, repo_path)

    lines: list[str] = [f"# {repo_path.name}", ""]
    for key in sorted(tree.keys()):
        _render_tree(tree[key], key, 2, lines)

    md_content = "\n".join(lines)

    if to_docx:
        from docx import Document  # type: ignore

        doc = Document()
        heading_re = re.compile(r"^(#{1,6})\s+(.*)")
        for line in lines:
            m = heading_re.match(line)
            if m:
                level = len(m.group(1))
                doc.add_heading(m.group(2), level=min(level, 9))
            else:
                doc.add_paragraph(line)
        doc.save(output_path)
    else:
        output_path.write_text(md_content, encoding="utf-8")

    console.log(f"Encoded {len(files)} files to [bold]{output_path}[/bold]")


_MODES = ("overwrite", "append", "ignore", "block", "merge")


def _two_way_merge(existing: str, incoming: str) -> tuple[str, bool]:
    """Merge two strings line-by-line; mark conflicts with git-style markers."""
    if existing == incoming:
        return existing, False
    a = existing.splitlines(keepends=True)
    b = incoming.splitlines(keepends=True)
    matcher = difflib.SequenceMatcher(None, a, b, autojunk=False)
    result: list[str] = []
    has_conflicts = False
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            result.extend(a[i1:i2])
        else:
            has_conflicts = True
            result.append("<<<<<<< existing\n")
            result.extend(a[i1:i2])
            result.append("=======\n")
            result.extend(b[j1:j2])
            result.append(">>>>>>> incoming\n")
    return "".join(result), has_conflicts


def _parse_markdown(md_content: str) -> dict[str, str]:
    """Return {relative_path: file_content} from markdown produced by encode."""
    lines = md_content.splitlines()
    heading_re = re.compile(r"^(#{1,6})\s+(.*)")
    fence_re = re.compile(r"^(`{3,})")

    path_stack: list[tuple[int, str]] = []
    file_map: dict[str, str] = {}
    i = 0

    while i < len(lines):
        m = heading_re.match(lines[i])
        if not m:
            i += 1
            continue

        depth = len(m.group(1))
        name = m.group(2).strip()

        path_stack = [(d, n) for d, n in path_stack if d < depth]
        path_stack.append((depth, name))

        if depth == 1:
            i += 1
            continue

        j = i + 1
        while j < len(lines) and lines[j].strip() == "":
            j += 1

        if j >= len(lines) or heading_re.match(lines[j]):
            i += 1
            continue

        fm = fence_re.match(lines[j]) if j < len(lines) else None
        if fm:
            opening_fence = fm.group(1)
            content_lines: list[str] = []
            k = j + 1
            while k < len(lines):
                if lines[k].strip() == opening_fence:
                    break
                content_lines.append(lines[k])
                k += 1
            while content_lines and content_lines[0] == "":
                content_lines.pop(0)

            rel_path = "/".join(n for _, n in path_stack[1:])
            file_map[rel_path] = "\n".join(content_lines)
            i = k + 1
        else:
            i += 1

    return file_map


def decode(args: argparse.Namespace) -> None:
    input_path = Path(args.input)
    mode = args.mode

    if input_path.suffix.lower() == ".docx":
        from docx import Document  # type: ignore

        doc = Document(input_path)
        raw_lines: list[str] = []
        for para in doc.paragraphs:
            style = para.style.name
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                raw_lines.append("#" * level + " " + para.text)
            else:
                raw_lines.append(para.text)
        md_content = "\n".join(raw_lines)
    else:
        md_content = input_path.read_text(encoding="utf-8")

    h1_match = re.search(r"^# (.+)$", md_content, re.MULTILINE)
    if not h1_match:
        raise ValueError("No H1 heading found in input — cannot determine repo name.")
    repo_name = h1_match.group(1).strip()

    file_map = _parse_markdown(md_content)
    output_dir = Path(args.output).resolve() / repo_name
    output_dir.mkdir(parents=True, exist_ok=True)

    for rel_path, content in file_map.items():
        dest = output_dir / rel_path
        dest.parent.mkdir(parents=True, exist_ok=True)

        if dest.exists():
            if mode == "block":
                raise FileExistsError(f"[block] Conflict: {dest} already exists.")
            elif mode == "ignore":
                console.log(f"[yellow]ignore[/yellow]    {rel_path}")
                continue
            elif mode == "overwrite":
                dest.write_text(content, encoding="utf-8")
                console.log(f"[red]overwrite[/red]  {rel_path}")
            elif mode == "append":
                existing = dest.read_text(encoding="utf-8")
                dest.write_text(existing + "\n----\n" + content, encoding="utf-8")
                console.log(f"[blue]append[/blue]     {rel_path}")
            elif mode == "merge":
                existing = dest.read_text(encoding="utf-8")
                merged, has_conflicts = _two_way_merge(existing, content)
                dest.write_text(merged, encoding="utf-8")
                if has_conflicts:
                    console.log(f"[yellow]conflict[/yellow]   {rel_path}")
                else:
                    console.log(f"[green]merged[/green]     {rel_path}")
        else:
            dest.write_text(content, encoding="utf-8")
            console.log(f"[green]create[/green]     {rel_path}")

    console.log(f"Decoded {len(file_map)} files to [bold]{output_dir}[/bold]")


def main() -> None:
    parser = argparse.ArgumentParser(description="Compress/decompress a git repo to/from markdown.")
    sub = parser.add_subparsers(dest="command", required=True)

    enc = sub.add_parser("encode", help="Compress a git repository to a markdown or docx file.")
    enc.add_argument("repo", help="Local path to the git repository.")
    enc.add_argument("-o", "--output", default="output.md", help="Output file (default: output.md). Use .docx for Word format.")

    dec = sub.add_parser("decode", help="Decompress a markdown or docx file to a directory.")
    dec.add_argument("input", help="Input markdown or docx file.")
    dec.add_argument("-o", "--output", default=".", help="Output directory (default: current directory).")
    dec.add_argument("-m", "--mode", choices=_MODES, default="block", help="Conflict resolution mode (default: block). 'merge' attempts a line-level merge and writes git conflict markers on conflicts.")

    args = parser.parse_args()
    if args.command == "encode":
        encode(args)
    else:
        decode(args)


if __name__ == "__main__":
    main()
